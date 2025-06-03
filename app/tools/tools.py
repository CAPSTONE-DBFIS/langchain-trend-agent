import asyncio
import io
import json
import logging
import os
import re
import time
from datetime import datetime, timedelta, timezone
from functools import wraps
from typing import Dict, List, Union
from urllib.parse import quote
from uuid import uuid4

import aiohttp
import FinanceDataReader as fdr
import openai
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
import requests
import wikipedia
import yfinance as yf
from bs4 import BeautifulSoup
from dateutil import parser
from docx import Document
from docx.shared import Inches
from dotenv import load_dotenv
from fake_useragent import UserAgent
from googleapiclient.discovery import build
from langchain_community.chat_models import ChatOpenAI
from langchain.chains.llm import LLMChain
from langchain.prompts import PromptTemplate
from langchain.tools import tool
from langchain_community.tools import WikipediaQueryRun
from langchain_community.utilities import WikipediaAPIWrapper
from langchain_tavily import TavilySearch
from pytrends.request import TrendReq
from pypdf import PdfReader
from requests.auth import HTTPBasicAuth
from unidecode import unidecode
from zoneinfo import ZoneInfo
from twikit import Client

from app.utils.db_util import get_db_connection
from app.utils.redis_util import get_redis_client
from app.utils.s3_util import upload_chart_to_s3, get_s3_client_and_bucket
from app.utils.es_util import fetch_domestic_articles, fetch_foreign_articles, fetch_sentiment_distribution, \
    get_es_client
from app.tools.tools_schema import *

logger = logging.getLogger(__name__)
logging.basicConfig(level=logging.INFO)

def async_time_logger(name: str):
    def decorator(func):
        @wraps(func)
        async def wrapper(*args, **kwargs):
            start = time.perf_counter()
            result = await func(*args, **kwargs)
            end = time.perf_counter()
            logger.info(f"[{name}] Execution time: {(end - start):.3f} seconds")
            return result
        return wrapper
    return decorator

load_dotenv()
KST = timezone(timedelta(hours=9))

@tool(args_schema=DomesticNewsSearchSchema)
async def domestic_news_search_tool(
    keyword: str,
    start_date: str = None,
    end_date: str = None,
    articles_per_day: int = 3
) -> Dict[str, Any]:
    """
    Aggregates daily mention counts and articles for a keyword, generates a Plotly line chart, and uploads it to S3.

    Args:
        keyword (str): The search keyword.
        start_date (str, optional): Start date in YYYY-MM-DD format. Defaults to 30 days ago.
        end_date (str, optional): End date in YYYY-MM-DD format. Defaults to yesterday.
        articles_per_day (int, optional): Number of articles to fetch per day. Defaults to 3.

    Returns:
        Dict[str, Any]: Dictionary containing keyword, date range, dates, mention counts, articles, and chart URL.
    """
    # 날짜 기본값
    if start_date is None:
        start_date = (datetime.now(ZoneInfo("Asia/Seoul")) - timedelta(days=30)).strftime("%Y-%m-%d")
    if end_date is None:
        end_date = (datetime.now(ZoneInfo("Asia/Seoul")) - timedelta(days=1)).strftime("%Y-%m-%d")

    # Redis 캐시 조회
    r = get_redis_client()
    cache_key = f"news:es:agg_articles:{keyword}:{start_date}:{end_date}:{articles_per_day}"
    if (cached := r.get(cache_key)):
        return json.loads(cached)

    # Elasticsearch 집계 쿼리
    es = get_es_client()
    agg_query = {
        "query": {
            "bool": {
                "must": [
                    {"range": {"date": {"gte": f"{start_date}T00:00:00", "lte": f"{end_date}T23:59:59"}}},
                    {"multi_match": {"query": keyword, "fields": ["title^2", "content"]}}
                ]
            }
        },
        "aggs": {
            "by_date": {
                "date_histogram": {
                    "field": "date",
                    "calendar_interval": "day",
                    "format": "yyyy-MM-dd"
                }
            }
        },
        "size": 0
    }

    def get_article_query(date_str: str) -> Dict:
        return {
            "query": {
                "bool": {
                    "must": [
                        {"range": {"date": {"gte": f"{date_str}T00:00:00", "lte": f"{date_str}T23:59:59"}}},
                        {"multi_match": {"query": keyword, "fields": ["title^2", "content"]}}
                    ]
                }
            },
            "sort": [
                {"date": {"order": "desc"}},
                {"_score": {"order": "desc"}}
            ],
            "size": articles_per_day,
            "highlight": {
                "pre_tags": [""],
                "post_tags": [""],
                "fields": {"content": {"fragment_size": 500, "number_of_fragments": 3, "no_match_size": 500}}
            }
        }

    try:
        # 날짜별 언급량 집계
        agg_result = es.search(
            index=os.getenv("ELASTICSEARCH_DOMESTIC_INDEX_NAME"),
            body=agg_query
        )
        buckets = agg_result.get("aggregations", {}).get("by_date", {}).get("buckets", [])
        dates = [b["key_as_string"] for b in buckets]
        mention_counts = [b["doc_count"] for b in buckets]

        if not dates:
            return {"error": "해당 기간 내 기사 데이터가 없습니다."}

        # 날짜별 기사 목록
        articles_by_date: Dict[str, Any] = {}
        for date_str in dates:
            article_result = es.search(
                index=os.getenv("ELASTICSEARCH_DOMESTIC_INDEX_NAME"),
                body=get_article_query(date_str)
            )
            hits = article_result.get("hits", {}).get("hits", [])
            day_list = []
            for h in hits:
                src = h["_source"]
                snippet = h.get("highlight", {}).get("content", [src.get("content", "")[:1000]])[0]
                day_list.append({
                    "title": src.get("title", ""),
                    "content": snippet,
                    "date": src.get("date", ""),
                    "url": src.get("url", ""),
                    "media_company": src.get("media_company", "")
                })
            articles_by_date[date_str] = day_list

        chart_url = None

        # 최고점 계산
        max_count = max(mention_counts)
        max_idx = mention_counts.index(max_count)
        peak_date = dates[max_idx]
        peak_count = max_count

        # Plotly Figure 생성 – 날짜가 2개 이상일 때만
        if len(dates) > 1:
            fig = go.Figure()

            fig.add_trace(
                go.Scatter(
                    x=dates,
                    y=mention_counts,
                    mode="lines+markers",
                    name=f"'{keyword}' 언급량",
                    line=dict(color="red", width=2),
                    marker=dict(size=8, color="red", line=dict(width=2, color="#FFFFFF")),
                    hovertemplate="%{x}<br>언급량: %{y}<extra></extra>",
                    fill="tozeroy",
                    fillcolor="rgba(255, 0, 0, 0.2)"
                )
            )

            # 최고점 강조
            fig.add_trace(
                go.Scatter(
                    x=[peak_date],
                    y=[peak_count],
                    mode="markers",
                    name="최고점",
                    marker=dict(size=12, color="darkred", symbol="circle"),
                    hovertemplate=f"{peak_date}<br>최고 언급량: {peak_count}<extra></extra>"
                )
            )

            # 최고점 날짜 텍스트 표시
            fig.add_annotation(
                x=peak_date,
                y=peak_count,
                text=f"최고점: {peak_date}",  # “최고점: YYYY-MM-DD” 형태로 표시
                showarrow=True,
                arrowhead=2,
                ax=0,
                ay=-30,
                font=dict(color="darkred", size=12)
            )

            fig.update_layout(
                title=f"{start_date} ~ {end_date} '{keyword}' 일일 언급량 추이",
                xaxis=dict(
                    title="날짜",
                    tickangle=-45,
                    tickmode="auto",
                    nticks=8,
                    tickformat="%m월 %d일"
                ),
                yaxis=dict(
                    title="언급량",
                    tickmode="auto",
                    nticks=6,
                    rangemode="tozero"
                ),
                font=dict(family="Noto Sans CJK KR"),
                height=500,
                margin=dict(l=40, r=40, t=80, b=120)
            )

            # S3 업로드
            s3_key = (
                f"daily/{datetime.now(ZoneInfo('Asia/Seoul')).strftime('%Y-%m-%d')}/"
                f"news_daily_trend_{uuid4().hex[:6]}.png"
            )
            chart_url = upload_chart_to_s3(fig, s3_key)

        # 날짜별 언급량
        mentions_by_date = {date: count for date, count in zip(dates, mention_counts)}

        results_list: List[Dict[str, Any]] = []
        for date_str, day_list in articles_by_date.items():
            for art in day_list:
                results_list.append({
                    "title": art["title"],
                    "content": art["content"],
                    "url": art["url"]
                })

        result_data = {
            "keyword": keyword,
            "mentions_by_date": mentions_by_date,
            "articles": articles_by_date,
            "results": results_list,
            "chart_url": chart_url
        }

        r.setex(cache_key, timedelta(days=7), json.dumps(result_data, ensure_ascii=False))
        return result_data

    except Exception as e:
        return {"error": f"Elasticsearch search failed: {str(e)}"}


@tool(args_schema=ForeignNewsSearchSchema)
@async_time_logger("foreign_news_search_tool")
async def foreign_news_search_tool(
    en_keyword: str, lang: str = "en", country: str = "us", max_results: int = 10
) -> Dict[str, Any]:
    """
    Foreign News Search Tool Using GNEWS API

    When to use:
        - When searching for foreign news articles.
        - When analyzing global trends with English keywords.

    Args:
        en_keyword (str): English keyword for search
        lang (str): Language code, defaults to 'en'
        country (str): Country code, defaults to 'us'
        max_results (int): Maximum number of articles (default 10, max 20)

    Returns:
        Dict[str, Any]:
            - keyword (str): Search keyword
            - lang (str): Language code
            - country (str): Country code
            - results (List[Dict]): List of articles with title, content, date, url, media_company
    """
    api_key = os.getenv("GNEWS_API_KEY")
    if not api_key:
        return {"error": "GNews API key is not set."}

    en_keyword = en_keyword.strip()
    if re.search(r"[!?&=+/\- ]", en_keyword) and not (en_keyword.startswith('"') and en_keyword.endswith('"')):
        en_keyword = f'"{en_keyword}"'
    encoded_query = quote(en_keyword)

    max_results = min(max_results, 20)
    url = f"https://gnews.io/api/v4/search?q={encoded_query}&lang={lang}&country={country}&max={max_results}&apikey={api_key}"

    try:
        async with aiohttp.ClientSession() as session:
            async with session.get(url) as response:
                if response.status != 200:
                    error_text = await response.text()
                    raise Exception(f"GNews API error {response.status}: {error_text}")
                data = await response.json()

        articles = data.get("articles", [])
        if not articles:
            return {
                "keyword": en_keyword,
                "lang": lang,
                "country": country,
                "results": []
            }

        result = {
            "keyword": en_keyword,
            "lang": lang,
            "country": country,
            "results": []
        }

        for article in articles[:max_results]:
            published_at = parser.parse(article["publishedAt"]).astimezone(KST).strftime("%Y-%m-%d %H:%M")
            result["results"].append({
                "title": article.get("title", ""),
                "content": article.get("content") or article.get("description") or "",
                "date": published_at,
                "url": article.get("url", ""),
                "media_company": article["source"]["name"]
            })

        return result

    except Exception as e:
        return {"error": f"GNews search failed: {str(e)}"}


@tool(args_schema=TrendKeywordSchema)
@async_time_logger("trend_keyword_tool")
async def trend_keyword_tool(*, period: str, date: str) -> Dict[str, Any]:
    """
    IT News Keyword Trend Tool

    When to use:
        - When analyzing trending keywords in IT news for a specific period.
        - When retrieving keyword frequencies, sentiment analysis, and related articles.

    Args:
        period (str): Period of analysis ('daily', 'weekly', or 'monthly')
        date (str): Reference date (YYYY-MM-DD)

    Returns:
        Dict[str, Any]:
            - date (str): Reference date
            - main_chart_url (str): URL of the combined (frequency + sentiment) stacked bar chart
            - keywords (List[Dict]): List of keyword data, each containing:
                - keyword (str): Keyword
                - frequency (int): Frequency of the keyword
                - sentiment_percent (Dict): Sentiment distribution percentages
                - articles (List[Dict]): List of articles with title, content, date, url, media_company
    """

    r = get_redis_client()
    cache_key = f"{period}_trend:{date}"
    if (cached := r.get(cache_key)):
        return json.loads(cached)

    # 기간 계산
    if period == "daily":
        start_date = end_date = date
    elif period == "weekly":
        end_date = date
        start_date = (datetime.fromisoformat(date) - timedelta(days=7)).strftime("%Y-%m-%d")
    elif period == "monthly":
        end_date = date
        start_date = (datetime.fromisoformat(date) - timedelta(days=30)).strftime("%Y-%m-%d")
    else:
        return {"error": "Period must be 'daily', 'weekly', or 'monthly'."}

    # 상위 10개 키워드 + 빈도 조회
    conn = get_db_connection()
    cur = conn.cursor()
    if period == "daily":
        cur.execute("""
                SELECT keyword, SUM(frequency)
                FROM keyword_frequencies
                WHERE date = %s
                GROUP BY keyword
                ORDER BY SUM(frequency) DESC
                LIMIT 10
            """, (date,))
    else:
        cur.execute("""
                SELECT keyword, SUM(frequency)
                FROM keyword_frequencies
                WHERE date BETWEEN %s AND %s
                GROUP BY keyword
                ORDER BY SUM(frequency) DESC
                LIMIT 10
            """, (start_date, end_date))
    rows = cur.fetchall()  # [(kw1, freq1), (kw2, freq2), ...]
    cur.close()
    conn.close()

    if not rows:
        return {"date": date, "main_chart_url": None, "keywords": []}

    # 키워드/빈도 딕셔너리
    keywords = [r0 for r0, _ in rows]
    keyword_freqs = {r0: int(r1) for r0, r1 in rows}

    # 감정 비율 조회
    sentiment_rows = []
    for kw in keywords:
        sent = await fetch_sentiment_distribution(keyword=kw, start_date=start_date, end_date=end_date)
        pos_pct = int(sent.get("positive_percent", 0))
        neu_pct = int(sent.get("neutral_percent", 0))
        neg_pct = int(sent.get("negative_percent", 0))
        sentiment_rows.append({
            "keyword": kw,
            "positive_pct": pos_pct,
            "neutral_pct": neu_pct,
            "negative_pct": neg_pct
        })

    df_sent = pd.DataFrame(sentiment_rows)

    # 실제 스택 높이 계산: “총 빈도수 * (감정 비율 / 100)”
    positive_heights = [
        int(keyword_freqs[kw] * row["positive_pct"] / 100)
        for kw, row in zip(df_sent["keyword"], sentiment_rows)
    ]
    neutral_heights = [
        int(keyword_freqs[kw] * row["neutral_pct"] / 100)
        for kw, row in zip(df_sent["keyword"], sentiment_rows)
    ]
    negative_heights = [
        int(keyword_freqs[kw] * row["negative_pct"] / 100)
        for kw, row in zip(df_sent["keyword"], sentiment_rows)
    ]

    # Plotly – stacked bar 생성
    fig_main = go.Figure()

    fig_main.add_trace(go.Bar(
        x=df_sent["keyword"],
        y=positive_heights,
        name="긍정",
        marker_color="seagreen",
        hovertemplate="%{x}<br>긍정 수: %{y}<extra></extra>"
    ))

    fig_main.add_trace(go.Bar(
        x=df_sent["keyword"],
        y=neutral_heights,
        name="중립",
        marker_color="lightgray",
        hovertemplate="%{x}<br>중립 수: %{y}<extra></extra>"
    ))

    fig_main.add_trace(go.Bar(
        x=df_sent["keyword"],
        y=negative_heights,
        name="부정",
        marker_color="salmon",
        hovertemplate="%{x}<br>부정 수: %{y}<extra></extra>",
        text=[keyword_freqs[kw] for kw in df_sent["keyword"]],
        textposition="outside"
    ))

    # 레이아웃 세팅
    fig_main.update_layout(
        barmode="stack",
        title=f"{start_date} ~ {end_date} 주요 키워드 빈도수 및 감정 분포",
        xaxis=dict(
            title="키워드",
            tickangle=-45
        ),
        yaxis=dict(
            title="총 언급 수"
        ),
        legend=dict(
            orientation="h",
            yanchor="bottom",
            y=1.02,
            xanchor="right",
            x=1
        ),
        font=dict(family="Noto Sans CJK KR"),
        height=500,
        margin=dict(l=40, r=40, t=80, b=120)
    )

    # S3에 업로드
    key_main = f"{period}/{date}/freq_sentiment_stacked_{uuid4().hex[:6]}.png"
    main_chart_url = upload_chart_to_s3(fig_main, key_main)

    # 각 키워드별 기사 목록 가져오기
    results = []
    for kw in keywords:
        articles = await fetch_domestic_articles(keyword=kw, start_date=start_date, end_date=end_date)
        results.append({
            "keyword": kw,
            "frequency": keyword_freqs[kw],
            "sentiment_percent": {
                "positive": int(df_sent.loc[df_sent["keyword"] == kw, "positive_pct"].iloc[0]),
                "neutral": int(df_sent.loc[df_sent["keyword"] == kw, "neutral_pct"].iloc[0]),
                "negative": int(df_sent.loc[df_sent["keyword"] == kw, "negative_pct"].iloc[0])
            },
            "articles": articles
        })

    output = {
        "date": date,
        "main_chart_url": main_chart_url,
        "keywords": results
    }

    # 캐시에 저장 (7일)
    r.setex(cache_key, timedelta(days=7), json.dumps(output, ensure_ascii=False))
    return output


@tool(args_schema=TrendReportSchema)
async def trend_report_tool(start_date=None, end_date=None):
    """
        Global IT News Trend Report Generation Tool

        When to use:
            - When a structured IT industry trend report is needed for a specific time range.
            - When analyzing domestic and foreign keyword frequencies and summarizing relevant news articles.
            - When generating a downloadable document (DOCX) with visualizations.

        Args:
            start_date (str, optional): Start date in format YYYY-MM-DD. Defaults to 7 days ago.
            end_date (str, optional): End date in format YYYY-MM-DD. Defaults to yesterday.

        Returns:
            str: Presigned URL to download the generated DOCX report
    """

    # 캐시 확인
    cache_key = f"trend_report:{start_date}:{end_date}"
    r = get_redis_client()
    if cached := r.get(cache_key):
        return cached.decode() if isinstance(cached, bytes) else cached

    # DB에서 상위 10개 국내/해외 키워드 조회
    conn = get_db_connection()
    cur = conn.cursor()

    cur.execute("""
            SELECT keyword, SUM(frequency)
            FROM keyword_analysis
            WHERE date BETWEEN %s AND %s
            GROUP BY keyword
            ORDER BY SUM(frequency) DESC
            LIMIT 10
        """, (start_date, end_date))
    dom_rows = cur.fetchall()
    if not dom_rows:
        return [{"error": f"No domestic keywords found between {start_date} and {end_date}."}]

    cur.execute("""
            SELECT keyword, SUM(frequency)
            FROM foreign_keyword_analysis
            WHERE date BETWEEN %s AND %s
            GROUP BY keyword
            ORDER BY SUM(frequency) DESC
            LIMIT 10
        """, (start_date, end_date))
    for_rows = cur.fetchall()
    if not for_rows:
        return [{"error": f"No foreign keywords found between {start_date} and {end_date}."}]

    cur.close()
    conn.close()

    # 국내: 감정 분포 스택형 차트 생성 함수
    async def make_domestic_sentiment_chart(rows: List[tuple], title_prefix: str) -> str:
        """
        rows: List of tuples [(keyword, freq), ...]
        title_prefix: "국내"
        반환값: 로컬에 저장된 차트 파일 경로 (PNG)
        """
        # 키워드와 빈도 dictionary
        keywords = [row[0] for row in rows]
        keyword_freqs: Dict[str, int] = {row[0]: int(row[1]) for row in rows}

        # 감정 분포 비율 조회
        sentiment_list = []
        for kw in keywords:
            sent = await fetch_sentiment_distribution(kw, start_date, end_date)
            pos_pct = int(sent.get("positive_percent", 0))
            neu_pct = int(sent.get("neutral_percent", 0))
            neg_pct = int(sent.get("negative_percent", 0))
            sentiment_list.append({
                "keyword": kw,
                "positive_pct": pos_pct,
                "neutral_pct": neu_pct,
                "negative_pct": neg_pct
            })

        df_sent = pd.DataFrame(sentiment_list)

        # 스택형 높이 계산
        positive_counts = [
            int(keyword_freqs[kw] * row["positive_pct"] / 100)
            for kw, row in zip(df_sent["keyword"], sentiment_list)
        ]
        neutral_counts = [
            int(keyword_freqs[kw] * row["neutral_pct"] / 100)
            for kw, row in zip(df_sent["keyword"], sentiment_list)
        ]
        negative_counts = [
            int(keyword_freqs[kw] * row["negative_pct"] / 100)
            for kw, row in zip(df_sent["keyword"], sentiment_list)
        ]

        # Plotly 스택형 바 차트
        fig = go.Figure()
        fig.add_trace(go.Bar(
            x=df_sent["keyword"],
            y=positive_counts,
            name="긍정 건수",
            marker_color="seagreen",
            hovertemplate="%{x}<br>긍정: %{y}<extra></extra>"
        ))
        fig.add_trace(go.Bar(
            x=df_sent["keyword"],
            y=neutral_counts,
            name="중립 건수",
            marker_color="lightgray",
            hovertemplate="%{x}<br>중립: %{y}<extra></extra>"
        ))
        # 부정 trace에 전체 빈도 텍스트 표시
        fig.add_trace(go.Bar(
            x=df_sent["keyword"],
            y=negative_counts,
            name="부정 건수",
            marker_color="salmon",
            hovertemplate="%{x}<br>부정: %{y}<extra></extra>",
            text=[keyword_freqs[kw] for kw in df_sent["keyword"]],
            textposition="outside"
        ))

        fig.update_layout(
            barmode="stack",
            title=f"{start_date} ~ {end_date} {title_prefix} 키워드 빈도수 및 감정 분포",
            xaxis=dict(title="키워드", tickangle=-45),
            yaxis=dict(title="총 언급 수"),
            legend=dict(
                orientation="h",
                yanchor="bottom",
                y=1.02,
                xanchor="right",
                x=1
            ),
            font=dict(family="Noto Sans CJK KR"),
            height=450,
            margin=dict(l=40, r=40, t=60, b=100)
        )

        # 로컬에 이미지 저장
        base_dir = os.path.abspath("./data/reports")
        os.makedirs(base_dir, exist_ok=True)
        filename = f"{title_prefix}_sentiment_{start_date}_{end_date}_{uuid4().hex[:6]}.png"
        path = os.path.join(base_dir, filename)
        fig.write_image(path, format="png")
        return path

    # 해외: 단순 빈도 바 차트 생성 함수
    def make_foreign_bar_chart(rows: List[tuple], title_prefix: str) -> str:
        """
        rows: List of tuples [(keyword, freq), ...]
        title_prefix: "해외"
        반환값: 로컬에 저장된 차트 파일 경로 (PNG)
        """
        labels, freqs = zip(*rows)
        # Plotly 단순 바 차트 (빈도만)
        fig = go.Figure(go.Bar(
            x=list(labels),
            y=list(freqs),
            marker_color="skyblue",
            text=list(freqs),
            textposition="outside",
            hovertemplate="%{x}<br>빈도: %{y}<extra></extra>"
        ))
        fig.update_layout(
            title=f"{start_date} ~ {end_date} {title_prefix} 키워드 빈도",
            xaxis=dict(title="키워드", tickangle=-45),
            yaxis=dict(title="총 언급 수"),
            font=dict(family="Noto Sans CJK KR"),
            height=400,
            margin=dict(l=40, r=40, t=60, b=100)
        )

        # 로컬에 이미지 저장
        base_dir = os.path.abspath("./data/reports")
        os.makedirs(base_dir, exist_ok=True)
        filename = f"{title_prefix}_bar_{start_date}_{end_date}_{uuid4().hex[:6]}.png"
        path = os.path.join(base_dir, filename)
        fig.write_image(path, format="png")
        return path

    # 국내/해외 차트 생성
    dom_chart_path = await make_domestic_sentiment_chart(dom_rows, "국내")
    for_chart_path = make_foreign_bar_chart(for_rows, "해외")

    # 5) 기사 요약 부분 (기존 로직 유지)
    async def summarize_articles(article_list):
        summaries = []
        for article in article_list:
            summary = (
                f"- {article['title']} ({article['date']}, {article['media_company']})\n"
                f"  {article['content'].strip()}"
            )
            summaries.append(summary)
        return "\n".join(summaries)

    dom_articles = ""
    for kw, _ in dom_rows:
        articles = await fetch_domestic_articles(keyword=kw, start_date=start_date, end_date=end_date)
        dom_articles += f"[국내 키워드: {kw}]\n"
        dom_articles += await summarize_articles(articles) + "\n\n"

    for_articles = ""
    for kw, _ in for_rows:
        articles = await fetch_foreign_articles(keyword=kw, start_date=start_date, end_date=end_date)
        for_articles += f"[해외 키워드: {kw}]\n"
        for_articles += await summarize_articles(articles) + "\n\n"

    # LLM 호출
    prompt = PromptTemplate.from_template("""
    당신은 기업 리서치 보고서를 전문적으로 작성하는 AI입니다.
    다음은 국내외 IT 뉴스 키워드 빈도 분석 및 관련 기사 내용을 요약한 결과입니다.
    작성 시 다음 규칙을 반드시 따르세요:
    - 각 섹션은 다음과 같이 시작: "개요:", "국내 뉴스 분석:", "해외 뉴스 분석:", "결론:"
    - 섹션 제목은 반드시 맨 앞에 한 줄로 출력하세요 (예: "개요:")
    - markdown/기호 사용 없이 서술식 문단 작성
    - 분석에 활용된 키워드는 반드시 요약문에 녹여서 설명
    - 통계/수치가 포함될 경우 자연스럽게 설명에 포함

    [분석 기간]
    {start_date} ~ {end_date}

    [국내 키워드 요약]
    {domestic_keywords}

    [국내 뉴스 기사]
    {domestic_articles}

    [해외 키워드 요약]
    {foreign_keywords}

    [해외 뉴스 기사]
    {foreign_articles}
    """)

    chain = LLMChain(
        llm=ChatOpenAI(model="gpt-4.1", temperature=0.3),
        prompt=prompt
    )

    result = chain.invoke({
        "start_date": start_date,
        "end_date": end_date,
        "domestic_keywords": "\n".join([f"- {kw}: {freq}번" for kw, freq in dom_rows]),
        "domestic_articles": dom_articles,
        "foreign_keywords": "\n".join([f"- {kw}: {freq}번" for kw, freq in for_rows]),
        "foreign_articles": for_articles
    })
    gpt_text = result["text"]

    # LLM 응답에서 섹션별 본문 추출
    def extract_section(text, title):
        pattern = rf"{title}:\s*(.*?)(?=\n(?:개요|국내 뉴스 분석|해외 뉴스 분석|결론):|\Z)"
        match = re.search(pattern, text, re.DOTALL)
        return match.group(1).strip() if match else "[내용 없음]"

    # DOCX 작성
    doc = Document()
    doc.add_heading("개요", level=1)
    doc.add_paragraph(extract_section(gpt_text, "개요"))

    doc.add_heading("국내 뉴스 분석", level=1)
    doc.add_picture(dom_chart_path, width=Inches(5.5))
    doc.add_paragraph(extract_section(gpt_text, "국내 뉴스 분석"))

    doc.add_heading("해외 뉴스 분석", level=1)
    doc.add_picture(for_chart_path, width=Inches(5.5))
    doc.add_paragraph(extract_section(gpt_text, "해외 뉴스 분석"))

    doc.add_heading("결론", level=1)
    doc.add_paragraph(extract_section(gpt_text, "결론"))

    report_filename = f"TrenDB_Report_{start_date}_{end_date}_{uuid4().hex[:8]}.docx"
    report_path = os.path.join("./data/reports", report_filename)
    doc.save(report_path)

    # S3 업로드 및 캐싱
    s3_key = f"report/{os.path.basename(report_path)}"
    s3, bucket = get_s3_client_and_bucket()
    with open(report_path, "rb") as f:
        s3.put_object(
            Body=f,
            Bucket=bucket,
            Key=s3_key,
            ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    url = f"https://{bucket}.s3.amazonaws.com/{s3_key}"
    r.setex(cache_key, timedelta(days=7), url)
    return url


# 경쟁사 목록
COMPETITORS = [
    "삼성SDS", "LG CNS", "현대오토에버", "SK주식회사 C&C", "포스코DX",
    "두산디지털이노베이션", "롯데이노베이트", "CJ올리브네트웍스", "신세계아이앤씨",
    "현대IT&E", "농협정보시스템", "하나금융티아이", "아시아나IDT",
    "한진정보통신", "코오롱벤티트", "kt ds", "교보DTS"
]

@tool(args_schema=CompetitorAnalysisSchema)
async def competitor_analysis_tool(
    start_date: str,
    end_date: str
) -> Dict[str, Union[str, List[Dict[str, Union[str, int, float, List[Dict[str, str]]]]]]]:
    """
    Competitor Analysis Tool

    When to use:
        - When a consolidated competitor analysis is required for a specific time range.
        - When calculating competitor-wise mention frequency and sentiment distribution.
        - When generating a visualization (stacked sentiment bar chart) and retrieving up to three sample articles per competitor.

    Args:
        start_date (str): Analysis start date (YYYY-MM-DD).
        end_date   (str): Analysis end date (YYYY-MM-DD).

    Notes:
        - Aggregates data for the period [start_date] to [end_date].
        - Uses Elasticsearch filter and sub-aggregations to compute mention counts and sentiment counts per competitor.
        - Uploads chart to S3 and returns its presigned URL.
        - Retrieves up to three representative articles per competitor using fetch_domestic_articles.
    """

    # Elasticsearch 클라이언트 초기화
    es = get_es_client()
    index = os.getenv("ELASTICSEARCH_DOMESTIC_INDEX_NAME")
    if not index:
        return {"error": "환경 변수 ELASTICSEARCH_DOMESTIC_INDEX_NAME이 설정되어 있지 않습니다."}

    competitors_data: List[Dict[str, Any]] = []

    for comp in COMPETITORS:
        # Elasticsearch 언급량 집계
        try:
            resp_agg = es.search(
                index=index,
                body={
                    "size": 0,
                    "query": {
                        "bool": {
                            "must": [
                                {
                                    "bool": {
                                        "should": [
                                            {"match_phrase": {"title": comp}},
                                            {"match_phrase": {"content": comp}}
                                        ],
                                        "minimum_should_match": 1
                                    }
                                },
                                {
                                    "range": {
                                        "date": {
                                            "gte": f"{start_date}T00:00:00",
                                            "lte": f"{end_date}T23:59:59"
                                        }
                                    }
                                }
                            ]
                        }
                    }
                }
            )
            mention_count = resp_agg.get("hits", {}).get("total", {}).get("value", 0)
        except Exception:
            mention_count = 0

        # fetch_sentiment_distribution 호출 → 비율(%) 가져오기
        try:
            sent = await fetch_sentiment_distribution(
                keyword=comp,
                start_date=start_date,
                end_date=end_date,
                index=index
            )
            pos_pct = int(sent.get("positive_percent", 0))
            neu_pct = int(sent.get("neutral_percent", 0))
            neg_pct = int(sent.get("negative_percent", 0))
        except Exception:
            pos_pct = neu_pct = neg_pct = 0

        # 비율(%) × 언급량 → 절대 값(정수)
        positive_count = int(mention_count * pos_pct / 100)
        neutral_count = int(mention_count * neu_pct / 100)
        negative_count = int(mention_count * neg_pct / 100)

        # 대표 기사 5건 가져오기
        try:
            search_res = es.search(
                index=index,
                body={
                    "query": {
                        "bool": {
                            "must": [
                                {
                                    "multi_match": {
                                        "query": comp,
                                        "type": "phrase",
                                        "fields": ["title^2", "content"]
                                    }
                                },
                                {
                                    "range": {
                                        "date": {
                                            "gte": start_date,
                                            "lte": end_date
                                        }
                                    }
                                }
                            ]
                        }
                    },
                    "size": 5,
                    "sort": [{"date": {"order": "desc"}}]
                }
            )
            hits = search_res.get("hits", {}).get("hits", [])
            articles = []
            for hit in hits:
                source = hit["_source"]
                content_snippet = hit.get("highlight", {}).get("content", [source.get("content", "")[:500]])[0]
                articles.append({
                    "title": source.get("title", ""),
                    "content": content_snippet,
                    "date": source.get("date", ""),
                    "url": source.get("url", "")
                })
        except Exception:
            articles = []

        competitors_data.append({
            "name": comp,
            "article_count": mention_count,
            "positive_pct": pos_pct,
            "neutral_pct": neu_pct,
            "negative_pct": neg_pct,
            "positive_count": positive_count,
            "neutral_count": neutral_count,
            "negative_count": negative_count,
            "articles": articles
        })

    # 언급량이 0이거나 감성 건수 합계가 0인 경쟁사 제외
    filtered = [
        c for c in competitors_data
        if c["article_count"] > 0
           and (c["positive_count"] + c["neutral_count"] + c["negative_count"]) > 0
    ]

    if not filtered:
        # 경쟁사 언급이 하나도 없으면 바로 빈 결과 반환
        result = {
            "start_date": start_date,
            "end_date": end_date,
            "chart_url": None,
            "competitors": []
        }
        return result

    # 언급량 내림차순 정렬
    filtered.sort(key=lambda x: x["article_count"], reverse=True)

    # 스택 바 차트 그릴 데이터 준비
    names = [c["name"] for c in filtered]
    pos_vals = [c["positive_count"] for c in filtered]
    neu_vals = [c["neutral_count"] for c in filtered]
    neg_vals = [c["negative_count"] for c in filtered]
    total_vals = [c["article_count"] for c in filtered]

    # Plotly로 Stacked Bar Chart 생성
    fig = go.Figure()
    fig.add_trace(go.Bar(
        x=names,
        y=pos_vals,
        name="긍정 건수",
        marker_color="seagreen",
        text=pos_vals,
        textposition="inside",
        hovertemplate="%{x}<br>긍정: %{y}<extra></extra>"
    ))
    fig.add_trace(go.Bar(
        x=names,
        y=neu_vals,
        name="중립 건수",
        marker_color="lightgray",
        text=neu_vals,
        textposition="inside",
        hovertemplate="%{x}<br>중립: %{y}<extra></extra>"
    ))
    fig.add_trace(go.Bar(
        x=names,
        y=neg_vals,
        name="부정 건수",
        marker_color="salmon",
        text=neg_vals,
        textposition="inside",
        hovertemplate="%{x}<br>부정: %{y}<extra></extra>"
    ))

    fig.update_layout(
        barmode="stack",
        title=f"{start_date} ~ {end_date} 경쟁사 언급량 및 감성분포",
        xaxis=dict(
            title="경쟁사",
            tickangle=-45
        ),
        yaxis=dict(
            title="언급 건수",
            range=[0, max(total_vals) * 1.1]
        ),
        legend=dict(
            orientation="h",
            yanchor="bottom",
            y=1.02,
            xanchor="right",
            x=1
        ),
        font=dict(family="Noto Sans CJK KR"),
        height=500,
        margin=dict(l=40, r=40, t=80, b=120)
    )

    # S3 업로드
    combined_key = f"competitor/combined_{start_date}_{end_date}_{uuid4().hex[:6]}.png"
    combined_chart_url = upload_chart_to_s3(fig, combined_key)

    # 최종 응답
    result = {
        "start_date": start_date,
        "end_date": end_date,
        "chart_url": combined_chart_url,
        "competitors": filtered
    }

    # Redis 캐시 (7일)
    r = get_redis_client()
    cache_key = f"competitor_analysis:{start_date}:{end_date}"
    r.setex(cache_key, timedelta(days=7), json.dumps(result, ensure_ascii=False))

    return result

@tool(args_schema=GoogleTrendsSchema)
@async_time_logger("google_trends_tool")
async def google_trends_tool(query: str, start_date: str = None, end_date: str = None) -> Dict[str, Union[str, List[float], List[str]]]:
    """
    Google Trends Tool

    When to use:
        - When tracking changes in keyword interest over time.
        - When visualizing recent trend patterns.

    Args:
        query (str): Search keyword
        start_date (str, optional): Start date (YYYY-MM-DD), defaults to last month
        end_date (str, optional): End date (YYYY-MM-DD)

    Returns:
        Dict[str, Union[str, List[float], List[str]]]:
            - query (str): Searched keyword
            - interest_data (List[float]): Interest scores by date
            - dates (List[str]): Dates (YYYY-MM-DD)
            - chart_url (str): URL of the interest timeseries chart on S3
            - chart_description (str): Description of the chart
            - error (str, optional): Error message if applicable

    Notes:
        - chart_url and chart_description are used for visualization insertion.
    """
    try:
        pytrends = TrendReq(hl="ko", tz=540)

        # 날짜 범위 설정
        if start_date and end_date:
            timeframe = f"{start_date} {end_date}"
        else:
            timeframe = "today 1-m"

        # 요청
        pytrends.build_payload([query], cat=0, timeframe=timeframe, geo="KR", gprop="")

        trend_data = pytrends.interest_over_time()

        if trend_data is None or trend_data.empty:
            return {"error": f"No trending data found for '{query}'."}

        interest_data = trend_data[query].dropna().tolist()
        dates = trend_data.index.strftime('%Y-%m-%d').tolist()

        # 시계열 그래프 생성
        df_trend = pd.DataFrame({
            "date": dates,
            "interest": interest_data
        })

        fig = px.line(
            df_trend, x="date", y="interest",
            title=f"Google 트렌드 관심도 추이: {query}",
            markers=True
        )
        fig.update_layout(
            xaxis_title="날짜",
            yaxis_title="관심도",
            height=400,
            font=dict(family="Noto Sans CJK KR")
        )

        # S3에 업로드
        key = f"google_trends/{slugify(query)}_trend.png"
        chart_url = upload_chart_to_s3(fig, key)

        return {
            "query": query,
            "interest_data": interest_data,
            "dates": dates,
            "chart_url": chart_url,
            "chart_description": f"Google Trends interest timeseries chart for {query}"
        }

    except Exception as e:
        if '429' in str(e):
            return {"error": f"Rate limit exceeded for query '{query}'. Try again later."}
        return {"error": f"Error retrieving Google Trends data: {str(e)}"}

def slugify(text: str) -> str:
    """Convert Korean keywords to ASCII strings."""
    ascii_text = unidecode(text)
    ascii_text = ascii_text.lower()
    # 영숫자와 하이픈만 허용, 나머지는 언더바로 대체
    ascii_text = ''.join(c if c.isalnum() or c == '-' else '_' for c in ascii_text)
    return ascii_text.strip('_')


@async_time_logger("search_daum_blogs")
async def search_daum_blogs(keyword: str, max_results: int = 10) -> List[Dict[str, str]]:
    """
    Daum Blog Post Search Function

    Args:
        keyword (str): Search keyword
        max_results (int): Maximum number of results

    Returns:
        List[Dict[str, str]]: List of blog posts with title, url, content, datetime, source
    """
    headers = {"Authorization": f"KakaoAK {os.getenv('DAUM_API_KEY')}"}
    params = {"query": keyword, "size": max_results, "sort": "accuracy"}

    try:
        response = requests.get(os.getenv("DAUM_API_URL"), headers=headers, params=params)
        response.raise_for_status()
        data = response.json()

        return [
            {
                "title": item["title"],
                "url": item["url"],
                "content": item["contents"],
                "datetime": parser.parse(item["datetime"]).strftime("%Y-%m-%d %H:%M"),
                "source": "daum_blog"
            }
            for item in data.get("documents", [])
        ]
    except Exception as e:
        raise RuntimeError(f"Daum API error: {str(e)}")

def clean_html(text: str) -> str:
    """Remove HTML tags and entities"""
    if not text:
        return ""
    text = re.sub(r"<[^>]+>", "", text)  # Remove HTML tags
    text = re.sub(r"&[^;]*;", "", text)  # Remove HTML entities
    return text

@async_time_logger("search_naver_blogs")
async def search_naver_blogs(keyword: str, max_result: int = 10) -> List[Dict[str, str]]:
    """
    Naver Blog Search Function

    Args:
        keyword (str): Search keyword
        max_result (int): Maximum number of results

    Returns:
        List[Dict[str, str]]: List of blog posts with title, url, content, datetime, source
    """
    posts = []

    try:
        display = min(max_result, 100)
        url = f"{os.getenv('NAVER_API_URL')}?query={keyword}&display={display}"
        headers = {
            "X-Naver-Client-Id": os.getenv("NAVER_CLIENT_ID"),
            "X-Naver-Client-Secret": os.getenv("NAVER_CLIENT_SECRET")
        }

        response = requests.get(url, headers=headers)
        response.raise_for_status()
        data = response.json()

        for item in data.get("items", [])[:max_result]:
            post_date = datetime.strptime(item["postdate"], "%Y%m%d")
            posts.append({
                "title": clean_html(item["title"]),
                "url": item["link"],
                "content": clean_html(item["description"]),
                "datetime": post_date.strftime("%Y-%m-%d 00:00"),
                "source": "naver_blog"
            })
        return posts
    except Exception as e:
        raise RuntimeError(f"Naver blog search error: {str(e)}")


def get_reddit_access_token():
    """
    Obtain Reddit access token
    """
    auth = HTTPBasicAuth(os.getenv("REDDIT_CLIENT_ID"), os.getenv("REDDIT_CLIENT_SECRET"))
    headers = {
        "User-Agent": "web:com.dbfis.chatbot:v1.0.0 (by /u/Hot_Mission1860)",
        "Content-Type": "application/x-www-form-urlencoded"
    }
    data = {
        "grant_type": "password",
        "username": os.getenv("REDDIT_USERNAME"),
        "password": os.getenv("REDDIT_PASSWORD")
    }

    response = requests.post("https://www.reddit.com/api/v1/access_token", headers=headers, auth=auth, data=data)

    if response.status_code != 200:
        raise Exception(f"Reddit OAuth authentication failed: {response.json()}")

    return response.json().get("access_token")


@async_time_logger("search_reddit_posts")
async def search_reddit_posts(keyword: str, max_result: int = 10) -> List[Dict[str, str]]:
    """
    Reddit Post Search Function
    """
    try:
        access_token = get_reddit_access_token()
        headers = {
            "Authorization": f"bearer {access_token}",
            "User-Agent": "web:com.dbfis.chatbot:v1.0.0"
        }
        url = f"https://oauth.reddit.com/search?q={keyword}&limit={max_result}&sort=hot"
        response = requests.get(url, headers=headers)
        if response.status_code != 200:
            raise Exception(f"Reddit search failed: {response.json()}")
        data = response.json()
        return [
            {
                "title": item["data"]["title"],
                "url": f"https://www.reddit.com{item['data']['permalink']}",
                "content": item["data"].get("selftext", "").strip()[:500],
                "datetime": datetime.utcfromtimestamp(item["data"]["created_utc"]).replace(tzinfo=timezone.utc).astimezone(KST).strftime('%Y-%m-%d %H:%M'),
                "source": "reddit"
            }
            for item in data.get("data", {}).get("children", [])
        ]
    except Exception as e:
        raise RuntimeError(f"Reddit search error: {str(e)}")


@async_time_logger("search_x_tweets")
async def search_x_tweets(keyword: str, max_results: int = 10, min_faves: int = 10) -> List[Dict[str, Any]]:
    logger.info(f"[X] search_x_tweets 호출됨 | keyword: {keyword}, max_results: {max_results}, min_faves: {min_faves}")

    KST = timezone(timedelta(hours=9))
    SPAM_WORDS = ["무료", "프로모션", "클릭", "구독", "이벤트", "광고", "free", "click", "promo", "win"]

    load_dotenv()
    username = os.getenv("X_USERNAME")
    email = os.getenv("X_EMAIL")
    password = os.getenv("X_PASSWORD")
    cookies_file = "twitter_cookies.json"

    if not all([username, email, password]):
        logger.error("[X] 인증 정보 누락: .env 확인")
        return [{"source": "x", "error": "X_USERNAME, X_EMAIL, 또는 X_PASSWORD 환경 변수가 설정되지 않았습니다."}]

    max_results = max(10, min(max_results, 100))

    try:
        client = Client(language="ko-KR")

        if os.path.exists(cookies_file):
            try:
                client.load_cookies(cookies_file)
                logger.info("[X] 쿠키 로드 완료, 유효성 검사 중...")
                await client.get_user_by_screen_name(username)
            except Exception as e:
                logger.warning(f"[X] 쿠키 유효하지 않음. 재로그인 시도: {e}")
                await client.login(auth_info_1=username, auth_info_2=email, password=password)
                await client.get_user_by_screen_name(username)
                client.save_cookies(cookies_file)
        else:
            logger.info("[X] 쿠키 없음. 로그인 시도")
            await client.login(auth_info_1=username, auth_info_2=email, password=password)
            await client.get_user_by_screen_name(username)
            client.save_cookies(cookies_file)

        query = f"{keyword} -filter:replies" if min_faves == 0 else f"{keyword} -filter:replies min_faves:{min_faves}"
        logger.info(f"[X] 트윗 검색 쿼리: {query}")

        tweets = await client.search_tweet(query=query, product="Latest")
        logger.info(f"[X] 원본 트윗 수: {len(tweets)}")

        tweet_list = []
        for tweet in tweets:
            text = tweet.text.strip()
            if len(text) <= 20: continue
            if any(spam_word.lower() in text.lower() for spam_word in SPAM_WORDS): continue
            if text.count("#") > 5: continue
            if tweet.user.followers_count < 100: continue

            tweet_list.append({
                "id": tweet.id,
                "text": text,
                "created_at": parser.parse(tweet.created_at),
                "url": f"https://x.com/i/status/{tweet.id}"
            })

        logger.info(f"[X] 필터링 후 트윗 수: {len(tweet_list)}")

        tweet_list.sort(key=lambda x: x["created_at"], reverse=True)

        results = [
            {
                "title": f"트윗 ID: {tweet['id']}",
                "url": tweet['url'],
                "content": tweet['text'][:500],
                "datetime": tweet['created_at'].astimezone(KST).strftime("%Y-%m-%d %H:%M"),
                "source": "x"
            }
            for tweet in tweet_list[:max_results]
        ]

        logger.info(f"[X] 최종 반환 결과 수: {len(results)}")

        if not results:
            logger.warning(f"[X] '{keyword}'에 대한 결과 없음")
            return []

        await asyncio.sleep(5)
        return results

    except Exception as e:
        logger.error(f"[X] 트윗 검색 오류: {str(e)}")
        return [{"source": "x", "error": f"트윗 검색 실패: {str(e)}"}]


@tool(args_schema=CommunitySearchSchema)
@async_time_logger("community_search_tool")
async def community_search_tool(
    korean_keyword: str,
    english_keyword: str,
    platform: str = "all",
    max_results: int = 20
) -> Dict[str, Any]:
    """
    Blog(Naver, Daum) and Community(Reddit, X) Post Search Tool

    When to use:
        - When analyzing public sentiment on blogs or Social Network platforms.

    Args:
        korean_keyword (str): Korean keyword for search
        english_keyword (str): English keyword for search
        platform (str): 'all', 'daum', 'naver', 'reddit', 'x'
        max_results (int): Maximum number of results (default: 20)

    Returns:
        Dict[str, Any]:
            - korean_keyword (str): Korean keyword
            - english_keyword (str): English keyword
            - platform (str): Platform used
            - results (List[Dict]): List of posts with title, url, content, datetime, source
            - errors (List[Dict]): List of errors from platforms
    """
    # 플랫폼별 max_results 제한 (최대 20개)
    platform_max_results = min(max_results // 2, 20)

    tasks = []
    if platform in ["all", "daum"]:
        tasks.append(search_daum_blogs(korean_keyword, platform_max_results))
    if platform in ["all", "naver"]:
        tasks.append(search_naver_blogs(korean_keyword, platform_max_results))
    if platform in ["all", "reddit"]:
        tasks.append(search_reddit_posts(english_keyword, platform_max_results))
    if platform in ["all", "x"]:
        tasks.append(search_x_tweets(english_keyword, platform_max_results))

    results_nested = await asyncio.gather(*tasks, return_exceptions=True)
    results = []
    errors = []

    for sublist in results_nested:
        if isinstance(sublist, Exception):
            errors.append({"source": "unknown", "error": str(sublist)})
            continue
        for item in sublist:
            if "error" in item:
                errors.append(item)
            else:
                results.append(item)

    # 최신순 정렬
    results_sorted = sorted(
        results,
        key=lambda x: x["datetime"] if isinstance(x["datetime"], datetime) else parser.parse(x["datetime"]),
        reverse=True
    )

    if platform != "all":
        return {
            "korean_keyword": korean_keyword,
            "english_keyword": english_keyword,
            "platform": platform,
            "results": results_sorted[:max_results],
            "errors": errors
        }

    # platform == "all"인 경우 균등 분배
    platforms = set(post["source"] for post in results_sorted)
    per_platform = max_results // max(len(platforms), 1)

    balanced_results = []
    for plat in platforms:
        posts = [post for post in results_sorted if post["source"] == plat][:per_platform]
        balanced_results.extend(posts)

    # 게시물 부족한 경우
    remaining = [post for post in results_sorted if post not in balanced_results]
    needed = max_results - len(balanced_results)
    balanced_results.extend(remaining[:needed])

    return {
        "korean_keyword": korean_keyword,
        "english_keyword": english_keyword,
        "platform": platform,
        "results": balanced_results[:max_results],
        "errors": errors
    }


@tool(args_schema=YoutubeVideoSchema)
@async_time_logger("youtube_video_tool")
async def youtube_video_tool(
    query: str,
    max_results: int = 5,
    order: str = "relevance"
):
    """
    YouTube Video Search Tool using YouTube Data API

    When to use:
        - When searching for recent or popular YouTube videos.
        - When exploring trends or topics through video content.
        - When thumbnails or video links are needed for reports.

    Args:
        query (str): Search keyword
        max_results (int, optional): Maximum number of results (1-10, default 5)
        order (str, optional): Sorting method: 'relevance', 'date', 'viewCount' etc.

    Returns:
        List[Dict]: List of videos with videoId, title, description, channelTitle, publishedAt, thumbnailUrl, url
    """
    youtube = build("youtube", "v3", developerKey=os.getenv("YOUTUBE_API_KEY"))

    search_response = youtube.search().list(
        q=query,
        part="snippet",
        type="video",
        maxResults=max_results,
        regionCode="KR",
        order=order
    ).execute()

    results = [
        {
            "videoId": item["id"]["videoId"],
            "title": item["snippet"]["title"],
            "description": item["snippet"]["description"],
            "channelTitle": item["snippet"]["channelTitle"],
            "publishedAt": item["snippet"]["publishedAt"],
            "thumbnailUrl": item["snippet"]["thumbnails"]["high"]["url"],
            "url": f"https://www.youtube.com/watch?v={item['id']['videoId']}"
        }
        for item in search_response["items"]
    ]
    return results


@tool(args_schema=SearchWebSchema)
@async_time_logger("web_search_tool")
async def web_search_tool(keyword: str, max_results: int = 10, include_images: bool = False) -> List[Dict[str, str]]:
    """
    Real-time Web Page Search Tool using Tavily API

    When to use:
        - When quickly exploring recent websites, blogs, or articles.
        - When structured JSON search results are required.
        - When integrating with a URL content extraction tool for detailed information.

    Args:
        keyword (str): Search keyword
        max_results (int, optional): Maximum number of results (1-20, default 10)
        include_images (bool, optional): Include images in search results

    Returns:
        List[Dict[str, str]]: List of search results with title, content, and URL
    """
    try:
        tavily_tool = TavilySearch(
            max_results=max_results,
            include_images=include_images
        )
        result = await tavily_tool.ainvoke({"query": keyword})
        return result
    except Exception as e:
        logger.error(f"Tavily search failed: {str(e)}")
        return []

@tool(args_schema=RequestUrlSchema)
@async_time_logger("request_url_tool")
async def request_url_tool(input_url: str) -> List[Dict[str, Any]]:
    """
    Web or PDF Content Extraction Tool

    When to use:
        - When extracting full text from a specific URL.
        - When raw text is needed for summarization, translation, or analysis.

    Args:
        input_url (str): Absolute URL of an HTTP(S) webpage or PDF file

    Returns:
        List[Dict[str, Any]]: List containing extracted content or error message
    """
    headers = {
        "User-Agent": (
            "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
            "AppleWebKit/537.36 (KHTML, like Gecko) "
            "Chrome/120.0.0.0 Safari/537.36"
        )
    }

    try:
        response = requests.get(input_url, headers=headers, timeout=10, verify=True)
        response.raise_for_status()

        if input_url.lower().endswith(".pdf"):
            text = ""
            with io.BytesIO(response.content) as f:
                pdf = PdfReader(f)
                for page in pdf.pages:
                    extracted = page.extract_text()
                    if extracted:
                        text += extracted + "\n"
            if not text.strip():
                return [{"error": "Unable to extract text from PDF."}]
            return [{"content": text.strip()}]
        else:
            soup = BeautifulSoup(response.text, "html.parser")
            if soup.body:
                text = soup.body.get_text()[:5000]
            else:
                text = soup.get_text()[:5000]
            text = re.sub(r"\s+", " ", text).strip()
            if not text or len(text) < 100:
                return [{"error": "No valid text found."}]
            if "example domain" in text.lower():
                return [{"error": "No valid text found."}]
            return [{"content": text}]

    except requests.exceptions.SSLError:
        return [{"error": "[Blocked] Invalid SSL certificate detected, site deemed insecure."}]
    except requests.RequestException as e:
        return [{"error": f"[Request Failed] {str(e)}"}]
    except Exception as e:
        return [{"error": f"[Processing Error] {str(e)}"}]


@tool(args_schema=WikipediaSchema)
@async_time_logger("wikipedia_tool")
async def wikipedia_tool(query: str) -> List[Dict[str, Any]]:
    """
    Wikipedia Summary Tool (Korean prioritized)

    When to use:
        - When formal and academic definitions are needed.
        - When English summaries are automatically provided if Korean is unavailable.

    Args:
        query (str): Search keyword

    Returns:
        List[Dict[str, Any]]: List containing summarized content or error message
    """
    # redis 캐싱
    r = get_redis_client()

    # 한국어 우선, 실패 시 영어
    for lang in ["ko", "en"]:
        try:
            cache_key = f"wiki:{lang}:{query}"
            cached = r.get(cache_key)
            if cached:
                return [{"content": cached}]

            # wiki_client로 wikipedia 모듈 전달
            api_wrapper = WikipediaAPIWrapper(
                wiki_client=wikipedia,       # wikipedia-api 클라이언트
                top_k_results=2,            # 최대 2개 문서
                doc_content_chars_max=1500, # 요약 최대 1500자
                lang=lang                   # 언어 설정
            )
            wikipedia_tool = WikipediaQueryRun(
                api_wrapper=api_wrapper,
            )
            result = wikipedia_tool.run(query)
            summaries = result.split("\n")[:10]  # 최대 10줄 요약
            content = f"Wikipedia Summary ({lang}):\n" + "\n".join(summaries) if summaries else f"No information found for '{query}' on Wikipedia ({lang})."
            r.set(cache_key, content)
            return [{"content": content}]
        except Exception as e:
            if lang == "en":  # 영어까지 실패 시
                return [{"error": f"Wikipedia search error: {str(e)}"}]
            continue


@tool(args_schema=NamuwikiSchema)
@async_time_logger("namuwiki_tool")
async def namuwiki_tool(keyword: str) -> List[Dict[str, Any]]:
    """
    Namuwiki Content Crawling and Summary Tool

    When to use:
        - When detailed Korean-based information is needed.
        - When seeking more informal and up-to-date information compared to Wikipedia.

    Args:
        keyword (str): Search keyword

    Returns:
        List[Dict[str, Any]]: List containing summarized content or error message
    """
    # redis 캐싱
    r = get_redis_client()
    cache_key = f"namuwiki:{keyword}"
    cached = r.get(cache_key)
    if cached:
        return [{"content": cached}]

    try:
        base_url = "https://namu.wiki"
        encoded_keyword = quote(keyword)
        headers = {"User-Agent": UserAgent().random}

        # 직접 url 접근
        direct_url = f"{base_url}/w/{encoded_keyword}"
        response = requests.get(direct_url, headers=headers, timeout=10)
        html = response.text

        # HTML 파싱
        soup = BeautifulSoup(html, "html.parser")
        all_divs = soup.find_all("div")

        # 제거할 키워드 + 정규식 패턴 정의
        irrelevant_keywords = [
            "CC BY-NC-SA", "namu.wiki", "umanle S.R.L",
            "Google Privacy Policy", "Términos de uso",
            "문서 가져오기", "최근 수정 시각", "틀", "분류:",
            "펼치기", "접기",
            "편집 요청", "편집 권한이 부족합니다", "ACL 탭", "도움말"
        ]
        heading_pattern = re.compile(r"^\d+(\s*\.\s*\d+)*\s*\.")

        extracted, seen = [], set()
        for div in all_divs:
            text = div.get_text(separator=" ", strip=True)
            if (
                text
                and len(text) > 40
                and text not in seen
                and not any(bad in text for bad in irrelevant_keywords)
                and not heading_pattern.match(text)
            ):
                extracted.append(text)
                seen.add(text)

        content = "\n\n".join(extracted[:30]) if extracted else "Content is empty."
        r.set(cache_key, content)
        return [{"content": content}]

    except Exception as e:
        return [{"error": f"[Error] Namuwiki request failed: {str(e)}"}]


@tool(args_schema=StockHistorySchema)
@async_time_logger("stock_history_tool")
async def stock_history_tool(
    symbol: str,
    start_date: str,
    end_date: str,
) -> Dict[str, Any]:
    """
    Unified Stock OHLCV Retrieval Tool (Global + Korean)

    Args:
        symbol (str): Stock ticker (e.g., 'AAPL') or 6-digit Korean stock code (e.g., '005930').
        start_date (str): Start date in 'YYYY-MM-DD' format.
        end_date (str): End date in 'YYYY-MM-DD' format.

    Returns:
        dict: {
            symbol (str): Queried stock symbol,
            history (list[dict]): List of daily OHLCV data (open, high, low, close, volume),
            info (dict): Additional metadata (only for global stocks),
            status (str): 'success' or 'failed',
            message (str|None): Error details if failed,
            chart_url (str): S3 URL to the generated stock chart,
            chart_description (str): Human-readable description of the chart
        }
    """
    def is_korean_symbol(sym: str) -> bool:
        return bool(re.fullmatch(r"\d{6}", sym))

    response = {
        "symbol": symbol,
        "history": [],
        "info": {},
        "status": "failed",
        "message": None
    }

    try:
        if is_korean_symbol(symbol):
            df = fdr.DataReader(symbol, start=start_date, end=end_date)
            info = {}
        else:
            ticker = yf.Ticker(symbol)
            df = ticker.history(
                start=start_date,
                end=end_date,
                interval="1d",
                auto_adjust=True,
                back_adjust=False
            )
            df = df.dropna(subset=["Close", "Volume"])
            df = df[df["Volume"] > 0]
            info = ticker.info or {}

        if df.empty:
            response["message"] = f"No valid trading records found for '{symbol}'."
            return response

        records = [{
            "date": idx.strftime("%Y-%m-%d"),
            "open": float(row["Open"]),
            "high": float(row["High"]),
            "low": float(row["Low"]),
            "close": float(row["Close"]),
            "volume": int(row["Volume"])
        } for idx, row in df.iterrows()]

        response.update({
            "history": records,
            "info": info,
            "status": "success"
        })

        df_vis = pd.DataFrame(records)
        df_vis["date"] = pd.to_datetime(df_vis["date"])

        fig = go.Figure()
        fig.add_trace(go.Scatter(
            x=df_vis["date"], y=df_vis["close"],
            name="Close Price", mode="lines+markers",
            line=dict(color="blue")
        ))
        fig.add_trace(go.Bar(
            x=df_vis["date"], y=df_vis["volume"],
            name="Volume", yaxis="y2",
            marker_color="lightgray", opacity=0.5
        ))

        fig.update_layout(
            title=f"{symbol} 주가 및 거래량",
            xaxis=dict(title="날짜"),
            yaxis=dict(title="종가"),
            yaxis2=dict(title="거래량", overlaying="y", side="right"),
            height=400,
            font=dict(family="Noto Sans CJK KR")
        )
        fig.update_xaxes(rangebreaks=[dict(bounds=["sat", "mon"])])

        key = f"stocks/{slugify(symbol)}_chart.png"
        chart_url = upload_chart_to_s3(fig, key)

        response.update({
            "chart_url": chart_url,
            "chart_description": f"{symbol} stock price and volume chart"
        })

        return response

    except Exception as e:
        response["message"] = f"Failed to retrieve stock data: {str(e)}"
        return response


@tool(args_schema=Dalle3ImageGenerationSchema)
@async_time_logger("dalle3_image_generation_tool")
async def dalle3_image_generation_tool(prompt: str) -> List[Dict[str, Any]]:
    """
    DALL·E 3 Image Generation Tool

    When to use:
        - When visualizations of trends or concepts are needed.

    Args:
        prompt (str): Image description (natural language)

    Returns:
        List[Dict]:
            - content (str): 'Image generated successfully'
            - url (str): Image URL
    """
    openai.api_key = os.getenv("DALLE_API_KEY")

    try:
        llm = ChatOpenAI(model_name="gpt-4o-mini", temperature=0.7)

        template = PromptTemplate.from_template(rf"""
            You are an expert prompt engineer for DALL·E 3.
            Rewrite the following description in **English**, adding vivid, concrete visual details.

            Instructions:
            - Describe the **scene**, **characters/objects**, and **background**.
            - Specify a **style** (e.g., photorealistic, watercolor, cyberpunk).
            - Add a **mood/atmosphere** and, if relevant, **lighting** (e.g., warm sunlight, neon glow).
            - If possible, suggest a **camera angle** (e.g., wide shot, close-up).

            Original prompt:
            "{prompt}"

            Output a detailed prompt of 50-100 words in English.
        """)

        formatted_prompt = template.format(prompt=prompt)
        enhanced_prompt = llm.invoke(formatted_prompt)

        dalle_response = openai.images.generate(
            model="dall-e-3",
            prompt=enhanced_prompt.content,
            size="1024x1024",
            quality="standard",
            n=1
        )
        return [{"content": "Image generated successfully", "url": dalle_response.data[0].url}]

    except Exception as e:
        print(f"DALL·E generation error: {str(e)}")
        return [{"error": f"Image generation failed: {str(e)}"}]

@tool(args_schema=PaperSearchSchema)
@async_time_logger("paper_search_tool")
async def paper_search_tool(
        query: str,
        max_results: int = 10,
        start_date: Optional[str] = None,
        end_date: Optional[str] = None,
        sort_by: str = "relevance"
) -> Dict[str, Any]:
    """
    OpenAlex Academic Paper Search Tool

    When to use:
        - When exploring academic papers across all disciplines (including arXiv, journals, etc.).
        - When retrieving paper titles, abstracts, URLs, authors, and metadata.

    Args:
        query (str): Search keyword
        max_results (int): Maximum number of papers to return (default 10, max 10)
        start_date (str, optional): Search start date (YYYY-MM-DD) (default 90 days ago)
        end_date (str, optional): Search end date (YYYY-MM-DD) (default today)
        sort_by (str): Sort by 'relevance' or 'date'

    Returns:
        Dict[str, Any]:
            - query (str): Search keyword
            - results (List[Dict]): List of papers with title, abstract, published_date, url, authors
            - error (str, optional): Error message if applicable
    """
    # 기본 날짜 설정
    if start_date is None:
        start_date = (datetime.today() - timedelta(days=90)).strftime("%Y-%m-%d")
    if end_date is None:
        end_date = datetime.today().strftime("%Y-%m-%d")

    # 캐시 키 생성
    cache_key = f"paper:openalex:{query}:{max_results}:{start_date or ''}:{end_date or ''}:{sort_by}"
    r = get_redis_client()
    if cached := r.get(cache_key):
        return json.loads(cached)

    # 정렬 및 필터 설정
    filter_params = [f"title.search:{query.replace(' ', '%20')}"]
    if start_date:
        filter_params.append(f"from_publication_date:{start_date}")
    if end_date:
        filter_params.append(f"to_publication_date:{end_date}")

    sort_param = "relevance_score:desc" if sort_by == "relevance" else "publication_date:desc"
    if max_results > 10:
        max_results = 10  # OpenAlex 페이지당 최대 25, 여기서는 10으로 제한

    # OpenAlex API URL
    base_url = "https://api.openalex.org/works"
    query_url = f"{base_url}?filter={','.join(filter_params)}&sort={sort_param}&per_page={max_results}"

    try:
        # API 요청
        response = requests.get(query_url)
        response.raise_for_status()
        data = response.json()

        # 결과 수집
        results = []
        for paper in data.get("results", []):
            published_date = paper.get("publication_date", "N/A")
            abstract = paper.get("abstract") or paper.get("abstract_inverted_index")
            if isinstance(abstract, dict):
                abstract = " ".join(word for word, _ in abstract.items())[:1000]
            elif not abstract:
                abstract = "No abstract available."

            results.append({
                "title": paper.get("title", "No title"),
                "abstract": abstract.strip()[:1000],
                "published_date": published_date,
                "url": paper.get("primary_location", {}).get("landing_page_url", paper.get("id")),
                "authors": [author["author"]["display_name"] for author in paper.get("authorships", [])]
            })

        # 결과가 없거나 필터링 후 비어 있는 경우
        if not results:
            result = {
                "query": query,
                "results": [],
                "message": "No papers found."
            }
        else:
            result = {
                "query": query,
                "results": results[:max_results]
            }

        # 캐시에 저장 (7일 TTL)
        r.setex(cache_key, timedelta(days=7), json.dumps(result, ensure_ascii=False))
        return result

    except Exception as e:
        logger.error(f"OpenAlex search failed: {str(e)}")
        return {"error": f"Paper search failed: {str(e)}"}


tools = [
    web_search_tool,
    domestic_news_search_tool,
    foreign_news_search_tool,
    competitor_analysis_tool,
    trend_report_tool,
    trend_keyword_tool,
    google_trends_tool,
    community_search_tool,
    youtube_video_tool,
    request_url_tool,
    wikipedia_tool,
    namuwiki_tool,
    stock_history_tool,
    dalle3_image_generation_tool,
    paper_search_tool
]